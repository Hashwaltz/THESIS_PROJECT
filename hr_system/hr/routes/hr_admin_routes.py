from flask import Blueprint, render_template, request, session, current_app, redirect, url_for, flash, jsonify, Response, send_file, make_response
from flask_login import login_required, current_user
from sqlalchemy.exc import IntegrityError
from datetime import datetime
from ..models.user import User
from ..models.hr_models import Employee, Attendance, Leave, Department, Position, LeaveType, EmploymentType
from ..forms import EmployeeForm, AttendanceForm, LeaveForm, DepartmentForm
from ..utils import admin_required, generate_employee_id, get_attendance_summary, get_current_month_range, load_excel_to_df, unlock_xlsx
from .. import db
from datetime import timedelta, datetime, date, time
from sqlalchemy.orm import joinedload
from collections import defaultdict
from hr_system.hr.functions import parse_date
import os
import csv
import shutil
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from flask import send_file
import io
from io import BytesIO
import pdfkit
from werkzeug.utils import secure_filename
from docx import Document
import pandas as pd
import re
import numpy as np
import uuid
from sqlalchemy import func, and_, case, cast, Date
import json
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.pagesizes import LETTER
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))  # For Filipino/Unicode chars
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../"))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "hr_static")

hr_admin_bp = Blueprint(
    'hr_admin',
    __name__,
    template_folder=TEMPLATE_DIR,
    static_folder=STATIC_DIR,
    static_url_path='/hr/static'
)


# ------------------------- Dashboard -------------------------
@hr_admin_bp.route('/dashboard')
@login_required
@admin_required
def hr_dashboard():
    today = datetime.now().date()
    total_employees = Employee.query.count()
    total_users = User.query.filter_by(active=True).count()
    total_inactive = User.query.filter_by(active=False).count()
    total_departments = Department.query.count()

    recent_employees = Employee.query.order_by(Employee.created_at.desc()).limit(5).all()
    recent_leaves = Leave.query.order_by(Leave.created_at.desc()).limit(5).all()

   
    # --- Graph Data (Current Month Data) ---
    start_date, end_date = get_current_month_range()
    
    monthly_dates = []
    monthly_present_counts = []
    monthly_absent_counts = []
    monthly_late_counts = []
    
    current_date = start_date
    while current_date <= end_date:
        records_on_date = Attendance.query.filter_by(date=current_date).all()
        monthly_dates.append(current_date.strftime("%b %d")) # e.g. Sep 01
        monthly_present_counts.append(len([r for r in records_on_date if r.status == "Present"]))
        monthly_absent_counts.append(len([r for r in records_on_date if r.status == "Absent"]))
        monthly_late_counts.append(len([r for r in records_on_date if r.status == "Late"]))
        current_date += timedelta(days=1)

    # Employees per department
    dept_data = db.session.query(
        Department.name, db.func.count(Employee.id)
    ).join(Employee, Employee.department_id == Department.id)\
     .group_by(Department.name).all()

    dept_labels = [d[0] for d in dept_data]
    dept_counts = [d[1] for d in dept_data]

    return render_template(
        'hr/admin/admin_dashboard.html',
        total_employees=total_employees,
        total_users=total_users,
        total_departments=total_departments,
        recent_employees=recent_employees,
        recent_leaves=recent_leaves,
        monthly_dates=monthly_dates,
        monthly_present_counts=monthly_present_counts,
        monthly_absent_counts=monthly_absent_counts,
        monthly_late_counts=monthly_late_counts,
        dept_labels=dept_labels,
        dept_counts=dept_counts,
        total_inactive=total_inactive
    )



# ------------------------- Employees -------------------------
@hr_admin_bp.route('/employees')
@login_required
@admin_required
def view_employees():
    search = request.args.get('search', '')
    department_id = request.args.get('department_id', '')
    employment_type_id = request.args.get('employment_type_id', '')  # ✅ new filter
    page = request.args.get('page', 1, type=int)

    # Base query with joins
    query = Employee.query.options(
        joinedload(Employee.department),
        joinedload(Employee.position),
        joinedload(Employee.employment_type)  # ✅ include employment type join
    )

    # Apply search filter
    if search:
        query = query.filter(
            Employee.first_name.ilike(f"%{search}%") |
            Employee.last_name.ilike(f"%{search}%") |
            Employee.email.ilike(f"%{search}%")
        )

    # Apply department filter
    if department_id:
        query = query.filter(Employee.department_id == int(department_id))

    # ✅ Apply employment type filter
    if employment_type_id:
        query = query.filter(Employee.employment_type_id == int(employment_type_id))

    # Order employees alphabetically
    query = query.order_by(Employee.last_name.asc(), Employee.first_name.asc())

    # Paginate
    employees = query.paginate(page=page, per_page=10)

    # Fetch all filter dropdown data
    departments = Department.query.order_by(Department.name.asc()).all()
    employment_types = EmploymentType.query.order_by(EmploymentType.name.asc()).all() 

    return render_template(
        'hr/admin/admin_view_employees.html',
        employees=employees,
        search=search,
        departments=departments,
        employment_types=employment_types,  
        selected_department=department_id,
        selected_employment_type=employment_type_id 
    )


@hr_admin_bp.route('/employees/add', methods=['GET', 'POST'])
@login_required
@admin_required
def add_employee():
    if request.method == 'POST':
        try:
            # Get department and generate unique employee ID
            department_id = request.form['department']
            new_employee_id = generate_employee_id(department_id)

            # Safely parse dates
            date_hired = parse_date(request.form['date_hired'], "Date Hired")
            date_of_birth = parse_date(request.form['date_of_birth'], "Date of Birth")

            if not date_hired or not date_of_birth:
                flash("Invalid date format!", "danger")
                return redirect(url_for('hr_admin.add_employee'))

            # Convert salary to float
            salary_str = request.form.get('salary', '').strip()
            try:
                salary = float(salary_str) if salary_str else 0.0
            except ValueError:
                flash("Invalid salary value!", "danger")
                return redirect(url_for('hr_admin.add_employee'))

            # --- 1. Create User ---
            default_password = "password123"
            user = User(
                email=request.form['email'],
                first_name=request.form['first_name'],
                last_name=request.form['last_name'],
                role="employee",
                password=default_password
            )
            db.session.add(user)
            db.session.flush()  # ensure user.id is available

            # --- 2. Create Employee linked to User ---
            employee = Employee(
                employee_id=new_employee_id,
                user_id=user.id,
                first_name=request.form['first_name'],
                last_name=request.form['last_name'],
                middle_name=request.form.get('middle_name'),
                email=request.form['email'],
                phone=request.form['phone'],
                address=request.form['address'],
                department_id=department_id,
                position_id=request.form['position'],
                employment_type_id=request.form['employment_type_id'],  # ✅ Added
                salary=salary,
                date_hired=date_hired,
                date_of_birth=date_of_birth,
                gender=request.form['gender'],
                marital_status=request.form['marital_status'],
                emergency_contact=request.form['emergency_contact'],
                active=True
            )
            db.session.add(employee)
            db.session.commit()

            flash("Employee and user account created successfully!", "success")
            return redirect(url_for('hr_admin.view_employees'))

        except IntegrityError:
            db.session.rollback()
            flash("Error: Employee or User already exists!", "danger")
            return redirect(url_for('hr_admin.add_employee'))
        except Exception as e:
            db.session.rollback()
            flash(f"Unexpected error: {str(e)}", "danger")
            return redirect(url_for('hr_admin.add_employee'))

    # GET request: render form
    departments = Department.query.all()
    positions = Position.query.all()
    employment_types = EmploymentType.query.all()  # ✅ Include for dropdown
    return render_template(
        'hr/admin/admin_add.html',
        departments=departments,
        positions=positions,
        employment_types=employment_types
    )



@hr_admin_bp.route('/employees/export')
@login_required
@admin_required
def export_employees_excel():
    
    employees = Employee.query.filter_by(active=True).all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Employees"

    headers = ['Employee ID', 'First Name', 'Last Name', 'Email', 'Department', 'Status']
    ws.append(headers)

    for emp in employees:
        ws.append([
            emp.employee_id,
            emp.first_name,
            emp.last_name,
            emp.email or '',
            emp.department.name if emp.department else '',
            'Active' if emp.active else 'Inactive'
        ])

    # 🔹 Auto-adjust column width based on max content length
    for col in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[col_letter].width = adjusted_width

    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="employees_report.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )




@hr_admin_bp.route("/employees/<int:employee_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    positions = Position.query.all()
    departments = Department.query.all()
    employment_types = EmploymentType.query.all()  # ✅ fetch employment types

    if request.method == "POST":
        try:
            employee.first_name = request.form.get("first_name")
            employee.last_name = request.form.get("last_name")
            employee.middle_name = request.form.get("middle_name")
            employee.email = request.form.get("email")
            employee.phone = request.form.get("phone")
            employee.address = request.form.get("address")

            dept_id = request.form.get("department")
            pos_id = request.form.get("position")
            emp_type_id = request.form.get("employment_type_id")  # ✅ new field

            employee.department_id = int(dept_id) if dept_id else None
            employee.position_id = int(pos_id) if pos_id else None
            employee.employment_type_id = int(emp_type_id) if emp_type_id else None  # ✅ assign relationship

            salary_val = request.form.get("salary")
            employee.salary = float(salary_val) if salary_val else None

            def parse_date(date_str):
                try:
                    return datetime.strptime(date_str, "%Y-%m-%d").date() if date_str else None
                except ValueError:
                    return None

            employee.date_hired = parse_date(request.form.get("date_hired"))
            employee.date_of_birth = parse_date(request.form.get("date_of_birth"))
            employee.gender = request.form.get("gender")
            employee.marital_status = request.form.get("marital_status")
            employee.emergency_contact = request.form.get("emergency_contact")

            status_val = request.form.get("status")
            employee.active = True if status_val == "1" else False

            employee.updated_at = datetime.utcnow()
            db.session.commit()

            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return {"status": "success", "message": "Employee updated successfully!"}, 200

            flash("Employee updated successfully!", "success")
            return redirect(url_for("hr_admin.view_employees"))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating employee {employee_id}: {e}")
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return {"status": "error", "message": "Error updating employee. Please try again."}, 500

            flash("Error updating employee. Please try again.", "error")

    return render_template(
        "hr/admin/admin_edit.html",
        employee=employee,
        positions=positions,
        departments=departments,
        employment_types=employment_types,  # ✅ pass to template
    )




@hr_admin_bp.route('/employees/<int:employee_id>/service_record')
@login_required
@admin_required
def export_service_record(employee_id):
    employee = Employee.query.get_or_404(employee_id)

    # Create Word document
    doc = Document()

    # ==============================
    # HEADER SECTION
    # ==============================
    title = doc.add_paragraph("S E R V I C E   R E C O R D")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    para1 = doc.add_paragraph("Republic of the Philippines")
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para1.runs[0].font.size = Pt(11)

    para2 = doc.add_paragraph("NORZAGARAY, REGION 3")
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.runs[0].bold = True
    para2.runs[0].font.size = Pt(11)

    doc.add_paragraph("")  # spacing

    # ==============================
    # EMPLOYEE INFORMATION
    # ==============================
    doc.add_paragraph(f"Name : {employee.last_name.upper()}, {employee.first_name.upper()} {employee.middle_name or ''}")
    birth_date = employee.date_of_birth.strftime('%B %d, %Y') if employee.date_of_birth else ''
    doc.add_paragraph(f"Date and place of birth : {birth_date}")
    doc.add_paragraph("(If married woman, give full maiden name)")
    doc.add_paragraph("(Date herein should be checked from birth or baptismal certificate or some other reliable documents)")
    doc.add_paragraph("B.P. Number: __________     TIN #: __________")
    doc.add_paragraph("")  # spacing

    # ==============================
    # CERTIFICATION STATEMENT
    # ==============================
    cert_text = (
        "This is to certify that the employee named hereunder actually rendered services "
        "in this Office as shown by the service record below, each line of which is supported "
        "by appointment and other papers actually issued by this Office and approved by the "
        "authorities concerned."
    )
    cert_para = doc.add_paragraph(cert_text)
    cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")  # spacing

    # ==============================
    # SERVICE RECORD TABLE
    # ==============================
    headers = [
        "From", "To", "Designation Status (1)", "Annual Salary (2)",
        "Station / Place of Assignment", "Branch (3)", "Leave(s) w/out Pay Date", "Cause"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text

    # Add data row (current employment)
    row = table.add_row().cells
    row[0].text = employee.date_hired.strftime('%b %d, %Y') if employee.date_hired else ''
    row[1].text = "Present"
    row[2].text = employee.position.name if employee.position else ''
    row[3].text = f"{employee.salary or ''}"
    row[4].text = employee.department.name if employee.department else ''
    row[5].text = ""
    row[6].text = ""
    row[7].text = ""

    doc.add_paragraph("")  # spacing

    # ==============================
    # FOOTER SECTION
    # ==============================
    footer = (
        "Issued on compliance with Executive Order No. 54 dated August 10, 1954, and in accordance "
        "with Circular No. 68 dated August 10, 1954 of the System."
    )
    footer_para = doc.add_paragraph(footer)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("CERTIFIED CORRECT:")
    doc.add_paragraph("FERNANDO DG. CRUZ")
    doc.add_paragraph("Acting MHRMO")
    doc.add_paragraph("Page 1 of 1")
    doc.add_paragraph(date.today().strftime("%A, %B %d, %Y"))

    # Adjust font size
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    # ==============================
    # SAVE TO BYTES AND RETURN
    # ==============================
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f"service_record_{employee.employee_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@hr_admin_bp.route('/users', methods=['GET'])
@login_required
@admin_required
def view_users():
    # Get query params
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '').strip()
    role_filter = request.args.get('role', '').strip()
    status_filter = request.args.get('status', '').strip()  

    # Base query
    query = User.query

    # Apply search filter
    if search:
        query = query.filter(
            (User.first_name.ilike(f"%{search}%")) |
            (User.last_name.ilike(f"%{search}%")) |
            (User.email.ilike(f"%{search}%"))
        )

    # Apply role filter
    if role_filter:
        query = query.filter(User.role == role_filter)

    # ✅ Apply status filter
    if status_filter == "active":
        query = query.filter(User.active.is_(True))
    elif status_filter == "inactive":
        query = query.filter(User.active.is_(False))

    # Paginate results
    users = query.order_by(User.id.asc()).paginate(page=page, per_page=10)

    # Roles for dropdown
    roles = ['admin', 'employee', 'dept_head', 'officer']

    return render_template(
        'hr/admin/admin_view_users.html',
        users=users,
        roles=roles,
        search=search,
        role_filter=role_filter,
        status_filter=status_filter  
    )

    
@hr_admin_bp.route("/user/<int:user_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_user(user_id):
    user = User.query.get_or_404(user_id)
    positions = Position.query.all()
    departments = Department.query.all()

    if request.method == "POST":
        try:
            # Map form fields
            user.email = request.form.get("email")
            user.first_name = request.form.get("first_name")
            user.last_name = request.form.get("last_name")
            user.role = request.form.get("role")
            user.active = request.form.get("status") == "1"

            db.session.commit()

            # If AJAX, return JSON for SweetAlert (use 'status' + 'message' shape)
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "success", "message": "User updated successfully!"})

            # Otherwise normal redirect with flash
            flash("User updated successfully!", "success")
            return redirect(url_for("hr_admin.view_users"))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating user {user_id}: {e}")

            # Return JSON error for AJAX
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "error", "message": "Error updating user. Please try again."}), 500

            flash("Error updating user. Please try again.", "error")

    # GET request → render template
    return render_template(
        "hr/admin/edit_user.html",
        user=user,
        positions=positions,
        departments=departments
    )


@hr_admin_bp.route('/attendance')
@login_required
@admin_required
def attendance():
    page = request.args.get('page', 1, type=int)
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    employee_filter = request.args.get('employee', '').strip()
    department_filter = request.args.get('department', '').strip()

    # Base query
    query = Attendance.query.join(Employee).join(Employee.department)

    # ✅ Date filters
    try:
        if start_date and not end_date:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            query = query.filter(Attendance.date == start_date_obj)

        elif end_date and not start_date:
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            query = query.filter(Attendance.date == end_date_obj)

        elif start_date and end_date:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()

            # Swap if user entered in reverse
            if end_date_obj < start_date_obj:
                start_date_obj, end_date_obj = end_date_obj, start_date_obj

            query = query.filter(
                and_(
                    Attendance.date >= start_date_obj,
                    Attendance.date <= end_date_obj
                )
            )
    except ValueError:
        pass

    # ✅ Employee and Department filters
    if employee_filter:
        query = query.filter(Attendance.employee_id == int(employee_filter))
    if department_filter:
        query = query.filter(Employee.department_id == int(department_filter))

    # Pagination
    attendances = query.order_by(Attendance.date.desc()).paginate(page=page, per_page=20, error_out=False)

    # Lists for dropdowns
    employees = Employee.query.filter_by(active=True).all()
    departments = Department.query.order_by(Department.name.asc()).all()

    return render_template(
        'hr/admin/admin_view_attendance.html',
        attendances=attendances,
        employees=employees,
        departments=departments,
        start_date=start_date,
        end_date=end_date,
        employee_filter=employee_filter,
        department_filter=department_filter
    )

# ----------------- CONFIG -----------------
ALLOWED_EXTENSIONS = {'xls', 'xlsx'}
UPLOAD_FOLDER = "uploads/attendance"

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



# ----------------- CONFIG -----------------
ALLOWED_EXTENSIONS = {'xls', 'xlsx'}
UPLOAD_FOLDER = "uploads/attendance"

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS




# ----------------- UPLOAD & PREVIEW -----------------
@hr_admin_bp.route('/add_attendance', methods=['GET', 'POST'])
@login_required
def add_attendance():
    preview_data = []

    if request.method == 'POST' and 'file' in request.files:
        file = request.files.get("file")
        if not file or not allowed_file(file.filename):
            flash("Please upload a valid Excel file (.xls or .xlsx).", "danger")
            return redirect(request.url)

        # Save uploaded file
        filename = secure_filename(file.filename)
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        filepath = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}_{filename}")
        file.save(filepath)

        try:
            df = pd.read_excel(filepath, header=None)
            records = []
            current_id, current_name, current_dept = None, None, None
            attendance_date = None

            for _, row in df.iterrows():
                line = " ".join(str(x) for x in row if str(x) != "nan").strip()
                if not line or "tabling date" in line.lower():
                    continue

                # Extract attendance date
                if "Attendance date:" in line:
                    attendance_date = line.split(":")[-1].strip()
                    continue

                # Extract employee info
                if "User ID" in line and "Name" in line:
                    uid_part = line.split("User ID:")[-1]
                    name_part = uid_part.split("Name:")
                    id_value = name_part[0].strip() if len(name_part) > 0 else None

                    if len(name_part) > 1:
                        name_dept_part = name_part[1].split("Department:")
                        name = name_dept_part[0].strip()
                        dept = name_dept_part[1].strip() if len(name_dept_part) > 1 else "Unknown"
                    else:
                        name, dept = "Unknown", "Unknown"

                    current_id, current_name, current_dept = id_value, name, dept
                    continue

                # Extract times
                if ":" in line:
                    times = line.split()
                    time_in = times[0] if len(times) > 0 else None
                    time_out = times[1] if len(times) > 1 else None
                    day_value = attendance_date or datetime.now().date().isoformat()

                    # Check DB for employee
                    emp_match = None
                    try:
                        emp_id_int = int(float(current_id))
                        emp_match = Employee.query.get(emp_id_int)
                    except:
                        emp_match = None

                    # If matched, use DB name, else leave as Excel
                    record_name = emp_match.get_full_name() if emp_match else current_name or "Unknown"
                    matched = True if emp_match else False

                    records.append({
                        "Employee ID": current_id,
                        "Name": record_name,
                        "Department": getattr(emp_match.department, 'name', 'N/A') if emp_match else "Unknown",
                        "Day": day_value,
                        "Time In": time_in if matched else None,
                        "Time Out": time_out if matched else None,
                        "Matched": matched
                    })

            # Include unmatched active employees from DB not in Excel
            db_employees = Employee.query.filter_by(active=True).all()
            excel_ids = {int(float(r["Employee ID"])) for r in records if r["Matched"]}

            for emp in db_employees:
                if emp.id not in excel_ids:
                    # Mark absent
                    records.append({
                        "Employee ID": emp.id,
                        "Name": emp.get_full_name(),
                        "Department": getattr(emp.department, 'name', 'N/A') if emp.department else 'N/A',
                        "Day": attendance_date or datetime.now().date().isoformat(),
                        "Time In": None,
                        "Time Out": None,
                        "Matched": False
                    })

            if not records:
                flash("No valid attendance records found. Please check the Excel format.", "danger")
                return redirect(request.url)

            session['import_attendance_preview'] = records
            preview_data = records
            flash("Preview loaded. Please confirm import.", "info")

        except Exception as e:
            flash(f"Error reading Excel file: {e}", "danger")
            return redirect(request.url)

    return render_template('hr/admin/admin_import_attendance.html', preview=preview_data)


# ----------------- CONFIRM IMPORT -----------------
@hr_admin_bp.route('/add_attendance/confirm', methods=['POST'])
@login_required
def confirm_import_attendance():
    import os
    records = session.get('import_attendance_preview', [])
    if not records:
        flash("No attendance records to import.", "danger")
        return redirect(url_for('hr_admin.add_attendance'))

    imported_count = 0

    for row in records:
        emp_id = row.get("Employee ID")
        try:
            emp_id_int = int(float(emp_id))
        except:
            continue

        emp = Employee.query.get(emp_id_int)
        if not emp:
            continue

        day = row.get("Day")
        if not day or "Tabling" in str(day):
            continue

        # Convert day string to date
        try:
            date_list = []
            if "~" in day:
                start_str, end_str = day.split("~")
                start_date = pd.to_datetime(start_str.strip(), errors='coerce').date()
                end_date = pd.to_datetime(end_str.strip(), errors='coerce').date()
                if start_date and end_date:
                    date_list = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]
            else:
                single_date = pd.to_datetime(day.strip(), errors='coerce').date()
                if single_date:
                    date_list = [single_date]
        except:
            continue

        time_in = row.get("Time In")
        time_out = row.get("Time Out")

        for att_date in date_list:
            # ✅ Skip if already exists
            existing = Attendance.query.filter_by(employee_id=emp.id, date=att_date).first()
            if existing:
                continue

            time_in_obj = pd.to_datetime(time_in, errors='coerce').time() if time_in else None
            time_out_obj = pd.to_datetime(time_out, errors='coerce').time() if time_out else None

            new_att = Attendance(
                employee_id=emp.id,
                date=att_date,
                time_in=time_in_obj,
                time_out=time_out_obj,
                status="Present" if time_in_obj else "Absent",
                remarks=""
            )
            db.session.add(new_att)
            imported_count += 1

    db.session.commit()
    session.pop('import_attendance_preview', None)

    # ✅ Cleanup uploaded files
    try:
        if os.path.exists(UPLOAD_FOLDER):
            for f in os.listdir(UPLOAD_FOLDER):
                os.remove(os.path.join(UPLOAD_FOLDER, f))
    except Exception as e:
        print(f"⚠️ Cleanup error: {e}")

    flash(f"✅ Successfully imported {imported_count} attendance record(s).", "success")
    return redirect(url_for('hr_admin.add_attendance'))

@hr_admin_bp.route('/attendance/<int:attendance_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required # Adjust as needed (e.g., @officer_required)
def edit_attendance(attendance_id):
    attendance_record = Attendance.query.get_or_404(attendance_id)
    form = AttendanceForm(obj=attendance_record)
    form.employee_id.choices = [(e.id, f"{e.employee_id} - {e.get_full_name()}") for e in Employee.query.filter_by(active=True).all()]

    if form.validate_on_submit():
        try:
            form.populate_obj(attendance_record) # Populates the object with form data
            attendance_record.updated_at = datetime.utcnow() # Add an update timestamp if you have one
            db.session.commit()
            flash('Attendance record updated successfully!', 'success')
            return redirect(url_for('hr_admin.attendance')) # Redirect to the main attendance view
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating attendance record {attendance_id}: {e}")
            flash('Error updating attendance record. Please try again.', 'error')
    
    # If it's a GET request or form validation failed
    return render_template('hr/admin/admin_edit_attend.html', form=form, attendance_record=attendance_record)


@hr_admin_bp.route('/attendance/<int:attendance_id>/delete', methods=['POST'])
@login_required
@admin_required # Adjust as needed (e.g., @officer_required)
def delete_attendance(attendance_id):
    attendance_record = Attendance.query.get_or_404(attendance_id)
    try:
        db.session.delete(attendance_record)
        db.session.commit()
        flash('Attendance record deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting attendance record {attendance_id}: {e}")
        flash('Error deleting attendance record. Please try again.', 'error')

    return redirect(url_for('hr_admin.attendance'))


# ------------------------- Leaves -------------------------
@hr_admin_bp.route('/leaves')
@login_required
@admin_required
def view_leaves():
    page = request.args.get('page', 1, type=int)
    status_filter = request.args.get('status', '')

    query = Leave.query
    if status_filter:
        query = query.filter_by(status=status_filter)

    leaves = query.order_by(Leave.created_at.desc()).paginate(page=page, per_page=20, error_out=False)
    return render_template('hr/admin/admin_view_leaves.html', leaves=leaves, status_filter=status_filter)

@hr_admin_bp.route('/review-leaves')
@login_required
@admin_required
def review_leaves():
    page = request.args.get('page', 1, type=int)
    status_filter = request.args.get('status', '')

    query = Leave.query

    # Apply filter if selected
    if status_filter:
        query = query.filter_by(status=status_filter)

    # ✅ Sort: Pending → Approved → Rejected → (None if any)
    query = query.order_by(
        db.case(
            (Leave.status == 'Pending', 0),
            (Leave.status == 'Approved', 1),
            (Leave.status == 'Rejected', 2),
            else_=3
        ),
        Leave.created_at.desc()
    )

    # Pagination
    leaves_paginated = query.paginate(page=page, per_page=20, error_out=False)
    
    return render_template(
        'hr/admin/review_leaves.html',
        leaves=leaves_paginated,
        status_filter=status_filter
    )



@hr_admin_bp.route('/leaves/<int:leave_id>/action', methods=['POST'])
@login_required
@admin_required
def leave_action(leave_id):
    leave = Leave.query.get_or_404(leave_id)
    action = request.form.get('action')  # 'Approved' or 'Rejected'
    comments = request.form.get('comments', '').strip()  # strip extra spaces

    leave.status = action
    leave.comments = comments if comments else None  # set None if empty
    leave.approved_by = current_user.id
    leave.approved_at = datetime.utcnow()

    try:
        db.session.commit()
        flash(f'Leave request {action.lower()} successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Leave action error: {e}")
        flash('Error updating leave request.', 'error')

    return redirect(url_for('hr_admin.view_leaves'))



@hr_admin_bp.route('/departments')
@login_required
@admin_required
def view_departments():
    page = request.args.get('page', 1, type=int)

    # Get paginated departments
    departments = Department.query.paginate(page=page, per_page=10)

    # Create a dictionary mapping department_id -> employee_count
    employee_counts = dict(
        db.session.query(
            Employee.department_id,
            func.count(Employee.id)
        ).group_by(Employee.department_id).all()
    )

    return render_template(
        'hr/admin/admin_view_departments.html',
        departments=departments,
        employee_counts=employee_counts
    )


@hr_admin_bp.route('/departments/add', methods=['GET', 'POST'])
@login_required
@admin_required
def add_department():
    form = DepartmentForm()
    # Populate department head choices
    form.head_id.choices = [(u.id, f"{u.first_name} {u.last_name}") 
                            for u in User.query.filter(User.role.in_(['admin','officer','dept_head'])).all()]

    if form.validate_on_submit():
        department = Department(
            name=form.name.data,
            description=form.description.data,
            head_id=form.head_id.data if form.head_id.data else None
        )
        try:
            db.session.add(department)
            db.session.commit()
            flash('Department added successfully!', 'success')
            return redirect(url_for('hr_admin.add_department'))  # redirect to the same page to trigger SweetAlert
        except Exception as e:
            db.session.rollback()
            flash(f'Error adding department: {str(e)}', 'error')

    return render_template('hr/admin/admin_add_dept.html', form=form)

    

@hr_admin_bp.route("/department/<int:department_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_department(department_id):
    department = Department.query.get_or_404(department_id)
    employees = Employee.query.filter_by(department_id=department_id).all()

    if request.method == "POST":
        try:
            department.name = request.form.get("name") or department.name
            head_id = request.form.get("dept_head")
            department.head_id = int(head_id) if head_id else department.head_id
            department.description = request.form.get("description") or department.description

            db.session.commit()

            # ✅ Return JSON if it's an AJAX request
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify(status="success", message="Department updated successfully!")

            # Otherwise, redirect normally
            flash("Department updated successfully!", "success")
            return redirect(url_for("hr_admin.view_departments"))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating department {department_id}: {e}")

            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify(status="error", message="Error updating department. Please try again.")

            flash("Error updating department. Please try again.", "error")

    # GET request
    return render_template(
        "hr/admin/edit_dept.html",
        department=department,
        employees=employees
    )

@hr_admin_bp.route('/hr/admin/positions')
@login_required
def view_positions():
    """Display paginated list of positions with employee count."""
    page = request.args.get('page', 1, type=int)
    per_page = 10

    positions = Position.query.order_by(Position.name.asc()).paginate(page=page, per_page=per_page)

    # Count employees in each position
    employee_counts = (
        db.session.query(Employee.position_id, db.func.count(Employee.id))
        .group_by(Employee.position_id)
        .all()
    )
    employee_counts = {pos_id: count for pos_id, count in employee_counts}

    return render_template(
        'hr/admin/admin_view_positions.html',
        positions=positions,
        employee_counts=employee_counts
    )


@hr_admin_bp.route("/hr/admin/add_position", methods=["GET", "POST"])
@login_required
@admin_required
def add_position():
    from hr_system.hr.models.hr_models import Position

    if request.method == "POST":
        name = request.form.get("name").strip()
        description = request.form.get("description").strip()

        if not name:
            flash("Position name is required.", "error")
            return redirect(url_for("hr_admin.add_position"))

        # Check for duplicate name
        existing_position = Position.query.filter_by(name=name).first()
        if existing_position:
            flash("A position with this name already exists.", "error")
            return redirect(url_for("hr_admin.add_position"))

        # Create and save new position
        new_position = Position(name=name, description=description)
        db.session.add(new_position)
        db.session.commit()

        flash(f"Position '{name}' added successfully!", "success")
        return redirect(url_for("hr_admin.view_positions"))  # You can adjust this target route

    return render_template("hr/admin/add_positions.html")

@hr_admin_bp.route("/position/<int:position_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_position(position_id):
    position = Position.query.get_or_404(position_id)
    departments = Department.query.all()

    if request.method == "POST":
        try:
            position.name = request.form.get("name") or position.name
            position.description = request.form.get("description") or position.description
            dept_id = request.form.get("department_id")
            position.department_id = int(dept_id) if dept_id else position.department_id

            db.session.commit()

            # For AJAX request
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "success", "message": "Position updated successfully!"})

            flash("Position updated successfully!", "success")
            return redirect(url_for("hr_admin.view_positions"))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating position {position_id}: {e}")

            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "error", "message": "Error updating position. Please try again."})

            flash("Error updating position. Please try again.", "error")

    # Render form
    return render_template(
        "hr/admin/edit_position.html",
        position=position,
        departments=departments
    )

# ------------------------- Reports -------------------------
@hr_admin_bp.route('/reports')
@login_required
@admin_required
def reports():
    # Get filters from query parameters
    report_type = request.args.get('report_type', 'attendance')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    page = request.args.get('page', 1, type=int)

    # Convert dates to datetime objects if provided
    try:
        start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date() if start_date else None
        end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date() if end_date else None
    except ValueError:
        flash("Invalid date format", "error")
        start_date_obj = end_date_obj = None

    # Fetch data based on report type
    if report_type == 'attendance':
        query = Attendance.query
        if start_date_obj:
            query = query.filter(Attendance.date >= start_date_obj)
        if end_date_obj:
            query = query.filter(Attendance.date <= end_date_obj)
        data = query.order_by(Attendance.date.desc()).paginate(page=page, per_page=20)

        employees = Employee.query.filter_by(active=True).all()  # for filter dropdown

    elif report_type == 'leaves':
        query = Leave.query
        if start_date_obj:
            query = query.filter(Leave.start_date >= start_date_obj)
        if end_date_obj:
            query = query.filter(Leave.end_date <= end_date_obj)
        data = query.order_by(Leave.start_date.desc()).paginate(page=page, per_page=20)
        employees = Employee.query.filter_by(active=True).all()

    elif report_type == 'payroll':
        # Example payroll: just employees with salary (you can expand later)
        query = Employee.query.filter(Employee.salary != None)
        data = query.paginate(page=page, per_page=20)
        employees = None

    else:
        flash("Invalid report type", "error")
        return redirect(url_for('hr_admin.reports'))

    return render_template(
        'hr/admin/admin_reports.html',
        data=data,
        report_type=report_type,
        start_date=start_date or '',
        end_date=end_date or '',
        employees=employees if report_type in ['attendance', 'leaves'] else []
    )

@hr_admin_bp.route('/hr/admin/attendance_reports')
@login_required
@admin_required
def attendance_reports():
    """Generate and view reports for attendance data with working hours."""
    # --- Get filter inputs ---
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    department_id = request.args.get('department_id', type=int)

    start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date() if start_date_str else None
    end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else None

    # --- Build base query ---
    query = (
        db.session.query(
            Employee.id.label('employee_id'),
            Employee.first_name,
            Employee.last_name,
            Department.name.label('department_name'),
            func.count(
                case((Attendance.status != 'Absent', 1))
            ).label('days_present'),
            func.sum(Attendance.working_hours).label('total_working_hours')
        )
        .join(Attendance, Attendance.employee_id == Employee.id)
        .outerjoin(Department, Department.id == Employee.department_id)
    )

    # --- Apply filters ---
    if start_date:
        query = query.filter(Attendance.date >= start_date)
    if end_date:
        query = query.filter(Attendance.date <= end_date)
    if department_id:
        query = query.filter(Employee.department_id == department_id)

    # --- Group by Employee ---
    query = query.group_by(Employee.id, Department.name)
    results = query.all()

    # --- Process results ---
    report_data = []
    for r in results:
        report_data.append({
            'employee_name': f"{r.first_name} {r.last_name}",
            'department_name': r.department_name or 'N/A',
            'days_present': r.days_present,
            'total_hours': round(r.total_working_hours or 0, 2),
            'average_hours': round((r.total_working_hours or 0) / r.days_present, 2) if r.days_present else 0
        })

    # --- Get departments for dropdown ---
    departments = Department.query.order_by(Department.name.asc()).all()

    return render_template(
        'hr/admin/attendance_reports.html',
        report_data=report_data,
        departments=departments,
        start_date=start_date_str,
        end_date=end_date_str,
        department_id=department_id,
    )
@hr_admin_bp.route('/attendance/reports/word')
@login_required
@admin_required
def attendance_report_word():
    """Generate a Word report for attendance within a date range and optional department."""
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    department_id = request.args.get('department_id', type=int)

    # --- Convert dates ---
    start_date_dt = datetime.strptime(start_date, '%Y-%m-%d').date() if start_date else None
    end_date_dt = datetime.strptime(end_date, '%Y-%m-%d').date() if end_date else None

    # --- Base employee query ---
    query = db.session.query(Employee).join(Department, isouter=True)
    if department_id:
        query = query.filter(Employee.department_id == department_id)
    employees = query.order_by(Employee.last_name.asc()).all()

    # --- Build report data ---
    report_data = []
    for emp in employees:
        attendances = Attendance.query.filter_by(employee_id=emp.id)

        if start_date_dt:
            attendances = attendances.filter(Attendance.date >= start_date_dt)
        if end_date_dt:
            attendances = attendances.filter(Attendance.date <= end_date_dt)

        records = attendances.all()

        # Count by status
        days_present = sum(1 for a in records if a.status == "Present")
        days_late = sum(1 for a in records if a.status == "Late")
        days_absent = sum(1 for a in records if a.status == "Absent")

        # Compute total working hours using your model's field
        total_hours = round(sum(a.working_hours or 0 for a in records), 2)

        report_data.append({
            "employee_name": f"{emp.last_name}, {emp.first_name} {emp.middle_name or ''}".strip(),
            "department_name": emp.department.name if emp.department else "N/A",
            "days_present": days_present,
            "days_late": days_late,
            "days_absent": days_absent,
            "total_hours": total_hours
        })

    # --- Create Word document ---
    doc = Document()
    doc.add_heading("Norzagaray College - Attendance Report", level=1)
    doc.add_paragraph(
        f"Date Range: {start_date or 'All'} to {end_date or 'All'}"
    )
    if department_id:
        dept_name = Department.query.get(department_id).name
        doc.add_paragraph(f"Department: {dept_name}")
    doc.add_paragraph("Generated on: " + datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph("")

    # --- Create table ---
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["Employee", "Department", "Days Present", "Late", "Absent", "Total Hours"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Fill table
    for row in report_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row["employee_name"]
        row_cells[1].text = row["department_name"]
        row_cells[2].text = str(row["days_present"])
        row_cells[3].text = str(row["days_late"])
        row_cells[4].text = str(row["days_absent"])
        row_cells[5].text = f"{row['total_hours']:.2f}"

    # --- Save to memory buffer ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # --- Send file for download ---
    filename = f"attendance_report_{start_date or 'all'}_to_{end_date or 'all'}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@hr_admin_bp.route('/hr/admin/leave_reports')
@login_required
def leave_reports():
    """Generate and view reports for leave data."""
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    department_id = request.args.get('department_id', type=int)
    leave_type_id = request.args.get('leave_type_id', type=int)

    # Convert dates (optional but safer)
    start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date() if start_date_str else None
    end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else None

    # Base query
    query = db.session.query(
        Employee.id.label('employee_id'),
        func.concat(Employee.first_name, ' ', Employee.last_name).label('employee_name'),
        Department.name.label('department_name'),
        func.sum(Leave.days_requested).label('total_leave_days'),  # ✅ FIXED
        func.count(Leave.id).label('leave_count')
    ).join(Department, Employee.department_id == Department.id) \
     .join(Leave, Leave.employee_id == Employee.id)

    # Apply filters dynamically
    if start_date:
        query = query.filter(Leave.start_date >= start_date)
    if end_date:
        query = query.filter(Leave.end_date <= end_date)
    if department_id:
        query = query.filter(Employee.department_id == department_id)
    if leave_type_id:
        query = query.filter(Leave.leave_type_id == leave_type_id)

    query = query.group_by(Employee.id, Department.name)

    results = query.all()

    # Dropdown options
    departments = Department.query.all()
    leave_types = LeaveType.query.all()

    return render_template(
        'hr/admin/leave_reports.html',
        results=results,
        departments=departments,
        leave_types=leave_types,
        selected_dept=department_id,
        selected_leave_type=leave_type_id,
        start_date=start_date_str,
        end_date=end_date_str
    )


@hr_admin_bp.route('/leave/reports/word')
@login_required
@admin_required
def leave_report_word():
    """Generate a Word report for leave records within a date range and optional filters."""
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    department_id = request.args.get('department_id', type=int)
    leave_type_id = request.args.get('leave_type_id', type=int)

    # Convert date strings
    start_date_dt = datetime.strptime(start_date, '%Y-%m-%d').date() if start_date else None
    end_date_dt = datetime.strptime(end_date, '%Y-%m-%d').date() if end_date else None

    # --- Base Query ---
    query = (
        db.session.query(
            Employee.id.label("employee_id"),
            func.concat(Employee.first_name, ' ', Employee.last_name).label("employee_name"),
            Department.name.label("department_name"),
            LeaveType.name.label("leave_type"),
            func.count(Leave.id).label("leave_count"),
            func.sum(Leave.days_requested).label("total_leave_days")
        )
        .join(Department, Employee.department_id == Department.id)
        .join(Leave, Leave.employee_id == Employee.id)
        .join(LeaveType, LeaveType.id == Leave.leave_type_id)
    )

    # Apply filters dynamically
    if start_date_dt:
        query = query.filter(Leave.start_date >= start_date_dt)
    if end_date_dt:
        query = query.filter(Leave.end_date <= end_date_dt)
    if department_id:
        query = query.filter(Employee.department_id == department_id)
    if leave_type_id:
        query = query.filter(Leave.leave_type_id == leave_type_id)

    # Group by employee and leave type
    query = query.group_by(Employee.id, Department.name, LeaveType.name)
    results = query.all()

    # --- Create Word document ---
    doc = Document()
    doc.add_heading("Norzagaray College - Leave Report", level=1)
    doc.add_paragraph(f"Date Range: {start_date or 'All'} to {end_date or 'All'}")

    if department_id:
        dept = Department.query.get(department_id)
        if dept:
            doc.add_paragraph(f"Department: {dept.name}")

    if leave_type_id:
        leave_type = LeaveType.query.get(leave_type_id)
        if leave_type:
            doc.add_paragraph(f"Leave Type: {leave_type.name}")

    doc.add_paragraph("Generated on: " + datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph("")

    # --- Create table ---
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["Employee", "Department", "Leave Type", "Leave Count", "Total Leave Days"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # --- Fill table with query results ---
    for row in results:
        row_cells = table.add_row().cells
        row_cells[0].text = row.employee_name
        row_cells[1].text = row.department_name or "N/A"
        row_cells[2].text = row.leave_type or "N/A"
        row_cells[3].text = str(row.leave_count or 0)
        row_cells[4].text = str(row.total_leave_days or 0)

    # --- Save to memory ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # --- Send as downloadable Word file ---
    filename = f"leave_report_{start_date or 'all'}_to_{end_date or 'all'}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ----------------- EDIT PASSWORD ROUTE -----------------
@hr_admin_bp.route('/edit_password', methods=['GET', 'POST'])
@login_required
def edit_password():
    if request.method == 'POST':
        new_password = request.form.get('password', '').strip()
        if not new_password:
            flash("⚠️ Password cannot be empty.", "warning")
            return redirect(url_for('hr_admin.edit_password'))

        # Update password directly (no hashing)
        current_user.password = new_password
        db.session.commit()

        flash("✅ Password successfully updated.", "success")
        return redirect(url_for('hr_admin.edit_password'))

    # GET request → show the form
    return render_template('hr/admin/edit_profile.html')
